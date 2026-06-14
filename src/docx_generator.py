import docx
import copy
import os
import json
import sys
import re

def set_cell_text_with_style(cell, text, template_cell):
    font_name = 'Times New Roman'
    font_size = docx.shared.Pt(11)  # Default table cell size
    font_bold = None
    font_italic = None
    font_color = None
    
    # Extract style from template cell if possible
    if template_cell.paragraphs:
        p_template = template_cell.paragraphs[0]
        if p_template.runs:
            run_template = p_template.runs[0]
            font_bold = run_template.font.bold
            font_italic = run_template.font.italic
            font_color = run_template.font.color.rgb if run_template.font.color else None
            if run_template.font.size:
                font_size = run_template.font.size
            
    # Clear runs in target cell's paragraphs
    for p in cell.paragraphs:
        for r in list(p.runs):
            r_el = r._r
            r_el.getparent().remove(r_el)
            
    if not cell.paragraphs:
        p = cell.add_paragraph()
    else:
        p = cell.paragraphs[0]
        
    p.text = ""
    p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
    
    # Clean text (remove double spaces and tabs)
    clean_text = re.sub(r' +', ' ', text.strip().replace('\t', ''))
    
    # Parse markdown bold/italic
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', clean_text)
    for part in parts:
        if not part:
            continue
        is_bold = False
        is_italic = False
        clean_part = part
        
        if part.startswith('**') and part.endswith('**'):
            is_bold = True
            clean_part = part[2:-2]
        elif part.startswith('*') and part.endswith('*'):
            is_italic = True
            clean_part = part[1:-1]
            
        run = p.add_run(clean_part)
        run.font.name = font_name
        run.font.size = font_size
        
        if is_bold or font_bold:
            run.font.bold = True
        if is_italic or font_italic:
            run.font.italic = True
        if font_color:
            run.font.color.rgb = font_color
            
        # Force Times New Roman for cyrillic and all scripts in XML
        rPr = run._r.get_or_add_rPr()
        rFonts = docx.oxml.shared.OxmlElement('w:rFonts')
        rFonts.set(docx.oxml.ns.qn('w:ascii'), font_name)
        rFonts.set(docx.oxml.ns.qn('w:hAnsi'), font_name)
        rFonts.set(docx.oxml.ns.qn('w:cs'), font_name)
        rFonts.set(docx.oxml.ns.qn('w:eastAsia'), font_name)
        rPr.append(rFonts)


def populate_annotation_table(table, report_data):
    # Table 0: annotation summary
    fields = [
        report_data.get('topic', ''),
        report_data.get('relevance', ''),
        report_data.get('problem', ''),
        report_data.get('goal_and_tasks', ''),
        report_data.get('audience', ''),
        report_data.get('essence', ''),
        report_data.get('effects', '')
    ]
    for idx, text in enumerate(fields):
        if idx < len(table.rows):
            cell = table.rows[idx].cells[1]
            set_cell_text_with_style(cell, text, cell)

def rebuild_roadmap_table(table, roadmap_data):
    if len(table.rows) < 4:
        return
        
    category_template = copy.deepcopy(table.rows[2]._tr)
    event_template = copy.deepcopy(table.rows[3]._tr)
    
    cat_cell_template = table.rows[2].cells[0]
    event_cell_templates = [table.rows[3].cells[c] for c in range(8)]
    
    tasks = {}
    for item in roadmap_data:
        num_str = item.get('num', '1.1').strip()
        prefix = "1"
        if '.' in num_str:
            parts = [p for p in num_str.split('.') if p]
            if parts:
                prefix = parts[0]
        else:
            prefix = num_str[0] if num_str else "1"
            
        if prefix not in tasks:
            tasks[prefix] = []
        tasks[prefix].append(item)
        
    for task_num in sorted(tasks.keys()):
        # Append Category Row
        new_cat_tr = copy.deepcopy(category_template)
        for cell in new_cat_tr.xpath('.//w:tc'):
            for p in cell.xpath('.//w:p'):
                for r in list(p.xpath('.//w:r')):
                    r.getparent().remove(r)
        table._tbl.append(new_cat_tr)
        cat_row = table.rows[len(table.rows) - 1]
        set_cell_text_with_style(cat_row.cells[0], f"Задача {task_num}", cat_cell_template)
        
        # Append Event Rows
        for item in tasks[task_num]:
            new_event_tr = copy.deepcopy(event_template)
            for cell in new_event_tr.xpath('.//w:tc'):
                for p in cell.xpath('.//w:p'):
                    for r in list(p.xpath('.//w:r')):
                        r.getparent().remove(r)
            table._tbl.append(new_event_tr)
            row = table.rows[len(table.rows) - 1]
            
            # Map fields to cells (Col 0, 1, 3, 5, 6, 7)
            set_cell_text_with_style(row.cells[0], item.get('num', ''), event_cell_templates[0])
            set_cell_text_with_style(row.cells[1], item.get('task_name', ''), event_cell_templates[1])
            set_cell_text_with_style(row.cells[3], item.get('start_date', ''), event_cell_templates[3])
            set_cell_text_with_style(row.cells[5], item.get('end_date', ''), event_cell_templates[5])
            set_cell_text_with_style(row.cells[6], item.get('responsible', ''), event_cell_templates[6])
            set_cell_text_with_style(row.cells[7], item.get('result', ''), event_cell_templates[7])
            
    # Delete original template rows 2 to 9
    for _ in range(8):
        if len(table.rows) > 2:
            table.rows[2]._tr.getparent().remove(table.rows[2]._tr)

def rebuild_risks_table(table, risks_data):
    if len(table.rows) < 3:
        return
        
    risk_template = copy.deepcopy(table.rows[2]._tr)
    risk_cell_templates = [table.rows[2].cells[c] for c in range(4)]
    
    for item in risks_data:
        new_tr = copy.deepcopy(risk_template)
        for cell in new_tr.xpath('.//w:tc'):
            for p in cell.xpath('.//w:p'):
                for r in list(p.xpath('.//w:r')):
                    r.getparent().remove(r)
        table._tbl.append(new_tr)
        row = table.rows[len(table.rows) - 1]
        
        set_cell_text_with_style(row.cells[0], item.get('num', ''), risk_cell_templates[0])
        set_cell_text_with_style(row.cells[1], item.get('risk_name', ''), risk_cell_templates[1])
        set_cell_text_with_style(row.cells[2], item.get('mitigation', ''), risk_cell_templates[2])
        set_cell_text_with_style(row.cells[3], item.get('plan_b', ''), risk_cell_templates[3])
        
    # Delete original template rows 2 and 3
    for _ in range(2):
        if len(table.rows) > 2:
            table.rows[2]._tr.getparent().remove(table.rows[2]._tr)

def rebuild_stakeholders_table(table, stakeholders_data):
    if len(table.rows) < 2:
        return
        
    stakeholder_template = copy.deepcopy(table.rows[1]._tr)
    cell_templates = [table.rows[1].cells[c] for c in range(3)]
    
    for item in stakeholders_data:
        new_tr = copy.deepcopy(stakeholder_template)
        for cell in new_tr.xpath('.//w:tc'):
            for p in cell.xpath('.//w:p'):
                for r in list(p.xpath('.//w:r')):
                    r.getparent().remove(r)
        table._tbl.append(new_tr)
        row = table.rows[len(table.rows) - 1]
        
        set_cell_text_with_style(row.cells[0], item.get('organization', ''), cell_templates[0])
        set_cell_text_with_style(row.cells[1], item.get('position', ''), cell_templates[1])
        set_cell_text_with_style(row.cells[2], item.get('expectations', ''), cell_templates[2])
        
    # Delete original template rows 1 to 4
    for _ in range(4):
        if len(table.rows) > 1:
            table.rows[1]._tr.getparent().remove(table.rows[1]._tr)

def rebuild_communication_table(table, comm_data):
    if len(table.rows) < 2:
        return
        
    comm_template = copy.deepcopy(table.rows[1]._tr)
    cell_templates = [table.rows[1].cells[c] for c in range(4)]
    
    for item in comm_data:
        new_tr = copy.deepcopy(comm_template)
        for cell in new_tr.xpath('.//w:tc'):
            for p in cell.xpath('.//w:p'):
                for r in list(p.xpath('.//w:r')):
                    r.getparent().remove(r)
        table._tbl.append(new_tr)
        row = table.rows[len(table.rows) - 1]
        
        set_cell_text_with_style(row.cells[0], item.get('stakeholder', ''), cell_templates[0])
        set_cell_text_with_style(row.cells[1], item.get('topic', ''), cell_templates[1])
        set_cell_text_with_style(row.cells[2], item.get('method', ''), cell_templates[2])
        set_cell_text_with_style(row.cells[3], item.get('frequency', ''), cell_templates[3])
        
    # Delete original template rows 1 to 4
    for _ in range(4):
        if len(table.rows) > 1:
            table.rows[1]._tr.getparent().remove(table.rows[1]._tr)

def parse_markdown_text_to_runs(p, text, font_name='Times New Roman', font_size_pt=12):
    """Парсит текст с разметкой markdown (**жирный**, *курсив*) и добавляет в абзац в виде runs."""
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if not part:
            continue
        
        is_bold = False
        is_italic = False
        clean_part = part
        
        if part.startswith('**') and part.endswith('**'):
            is_bold = True
            clean_part = part[2:-2]
        elif part.startswith('*') and part.endswith('*'):
            is_italic = True
            clean_part = part[1:-1]
            
        run = p.add_run(clean_part)
        run.font.name = font_name
        run.font.size = docx.shared.Pt(font_size_pt)
        if is_bold:
            run.font.bold = True
        if is_italic:
            run.font.italic = True
            
        # Force font XML
        rPr = run._r.get_or_add_rPr()
        rFonts = docx.oxml.shared.OxmlElement('w:rFonts')
        rFonts.set(docx.oxml.ns.qn('w:ascii'), font_name)
        rFonts.set(docx.oxml.ns.qn('w:hAnsi'), font_name)
        rFonts.set(docx.oxml.ns.qn('w:cs'), font_name)
        rFonts.set(docx.oxml.ns.qn('w:eastAsia'), font_name)
        rPr.append(rFonts)

def get_markdown_lines(text):
    """Разбивает текст на отдельные строки, очищая от лишних пробелов и пустых строк."""
    raw_lines = text.replace('\r', '').split('\n')
    clean_lines = []
    for line in raw_lines:
        line_str = line.strip()
        if line_str:
            # Очищаем двойные пробелы
            line_str = re.sub(r' +', ' ', line_str)
            # Очищаем табуляцию
            line_str = line_str.replace('\t', '')
            clean_lines.append(line_str)
    return clean_lines

def format_paragraph_staff_style(p, text, font_name='Times New Roman', font_size_pt=12):
    p.text = ""
    p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    
    clean_text = text.strip()
    
    is_bullet = False
    is_numbered = False
    bullet_symbol = "— "  # Стандартный длинный тире-дефис для списков в РФ
    
    # Проверка на списки-буллеты
    if clean_text.startswith(("- ", "* ", "• ")):
        is_bullet = True
        clean_text = clean_text[2:].strip()
    # Проверка на нумерованные списки
    elif re.match(r'^\d+[\.\)]\s+', clean_text):
        is_numbered = True
        
    if is_bullet:
        p.paragraph_format.left_indent = docx.shared.Cm(1.25)
        p.paragraph_format.first_line_indent = docx.shared.Cm(-0.5)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_before = docx.shared.Pt(0)
        p.paragraph_format.space_after = docx.shared.Pt(3)
        
        run_bullet = p.add_run(bullet_symbol)
        run_bullet.font.name = font_name
        run_bullet.font.size = docx.shared.Pt(font_size_pt)
        
        # Форсируем шрифт для маркера списка
        rPr = run_bullet._r.get_or_add_rPr()
        rFonts = docx.oxml.shared.OxmlElement('w:rFonts')
        rFonts.set(docx.oxml.ns.qn('w:ascii'), font_name)
        rFonts.set(docx.oxml.ns.qn('w:hAnsi'), font_name)
        rFonts.set(docx.oxml.ns.qn('w:cs'), font_name)
        rFonts.set(docx.oxml.ns.qn('w:eastAsia'), font_name)
        rPr.append(rFonts)
        
        parse_markdown_text_to_runs(p, clean_text, font_name, font_size_pt)
        
    elif is_numbered:
        p.paragraph_format.left_indent = docx.shared.Cm(1.25)
        p.paragraph_format.first_line_indent = docx.shared.Cm(-0.5)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_before = docx.shared.Pt(0)
        p.paragraph_format.space_after = docx.shared.Pt(3)
        
        match = re.match(r'^(\d+[\.\)])\s+(.*)', clean_text)
        if match:
            num_part = match.group(1) + " "
            rest_part = match.group(2)
            
            run_num = p.add_run(num_part)
            run_num.font.name = font_name
            run_num.font.size = docx.shared.Pt(font_size_pt)
            run_num.font.bold = True
            
            # Форсируем шрифт для номера
            rPr = run_num._r.get_or_add_rPr()
            rFonts = docx.oxml.shared.OxmlElement('w:rFonts')
            rFonts.set(docx.oxml.ns.qn('w:ascii'), font_name)
            rFonts.set(docx.oxml.ns.qn('w:hAnsi'), font_name)
            rFonts.set(docx.oxml.ns.qn('w:cs'), font_name)
            rFonts.set(docx.oxml.ns.qn('w:eastAsia'), font_name)
            rPr.append(rFonts)
            
            parse_markdown_text_to_runs(p, rest_part, font_name, font_size_pt)
        else:
            parse_markdown_text_to_runs(p, clean_text, font_name, font_size_pt)
    else:
        # Обычный абзац
        p.paragraph_format.first_line_indent = docx.shared.Cm(1.25)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_before = docx.shared.Pt(0)
        p.paragraph_format.space_after = docx.shared.Pt(6)
        
        parse_markdown_text_to_runs(p, clean_text, font_name, font_size_pt)

def replace_text_under_heading(doc, heading_text_prefix, new_text):
    for idx, p in enumerate(doc.paragraphs):
        if p.style.name.startswith("Heading 1") and p.text.strip().startswith(heading_text_prefix):
            current_idx = idx + 1
            # Delete empty paragraphs/list items immediately following it (until next Heading 1)
            while current_idx < len(doc.paragraphs) and not doc.paragraphs[current_idx].style.name.startswith("Heading 1"):
                p_to_del = doc.paragraphs[current_idx]
                p_to_del._element.getparent().remove(p_to_del._element)
                
            # Now insert the new text paragraphs
            paragraphs_text = get_markdown_lines(new_text)
                
            target_p = p
            for pt in paragraphs_text:
                new_p_el = docx.oxml.shared.OxmlElement('w:p')
                target_p._element.addnext(new_p_el)
                new_p = docx.text.paragraph.Paragraph(new_p_el, target_p._parent)
                new_p.style = 'Normal'
                
                format_paragraph_staff_style(new_p, pt)
                target_p = new_p
            return True
    return False

def replace_text_under_appendix(doc, heading_text_prefix, new_text):
    for idx, p in enumerate(doc.paragraphs):
        if p.style.name.startswith("Heading 1") and heading_text_prefix in p.text:
            current_idx = idx + 1
            # Delete empty paragraphs/list items immediately following it (until next Heading 1)
            while current_idx < len(doc.paragraphs) and not doc.paragraphs[current_idx].style.name.startswith("Heading 1"):
                p_to_del = doc.paragraphs[current_idx]
                p_to_del._element.getparent().remove(p_to_del._element)
                
            # Now insert the new text paragraphs
            paragraphs_text = get_markdown_lines(new_text)
                
            target_p = p
            for pt in paragraphs_text:
                new_p_el = docx.oxml.shared.OxmlElement('w:p')
                target_p._element.addnext(new_p_el)
                new_p = docx.text.paragraph.Paragraph(new_p_el, target_p._parent)
                new_p.style = 'Normal'
                
                format_paragraph_staff_style(new_p, pt)
                target_p = new_p
            return True
    return False


def insert_diagram_to_docx(doc, heading_text_prefix, image_path):
    """Вставляет сгенерированную схему/диаграмму в конец указанного раздела отчета."""
    import os
    import docx
    
    if not os.path.exists(image_path):
        print(f"[DOCX ДИАГРАММА] Файл {image_path} не найден. Пропускаем.")
        return False
        
    for idx, p in enumerate(doc.paragraphs):
        if heading_text_prefix in p.text and p.style.name.startswith("Heading"):
            # Ищем конец этого раздела (до следующего заголовка Heading)
            current_idx = idx + 1
            while current_idx < len(doc.paragraphs) and not doc.paragraphs[current_idx].style.name.startswith("Heading"):
                current_idx += 1
            
            # Вставляем новый абзац прямо перед следующим заголовком
            insert_pos = current_idx - 1
            if insert_pos < len(doc.paragraphs):
                target_p = doc.paragraphs[insert_pos]
                new_p_el = docx.oxml.shared.OxmlElement('w:p')
                target_p._element.addnext(new_p_el)
                new_p = docx.text.paragraph.Paragraph(new_p_el, target_p._parent)
                new_p.style = 'Normal'
                new_p.alignment = 1 # Center
                
                # Добавляем изображение
                run = new_p.add_run()
                run.add_picture(image_path, width=docx.shared.Inches(5.8))
                
                # Добавляем подпись под рисунком
                caption_p_el = docx.oxml.shared.OxmlElement('w:p')
                new_p._element.addnext(caption_p_el)
                caption_p = docx.text.paragraph.Paragraph(caption_p_el, target_p._parent)
                caption_p.style = 'Normal'
                caption_p.alignment = 1 # Center
                
                # В соответствии с ГОСТ подпись рисунка: Рисунок X — Название
                fig_num = "1" if "1" in heading_text_prefix else "2"
                fig_title = "Дерево проблем и причинно-следственных связей" if fig_num == "1" else "Архитектура Цифровой инвестиционной экосистемы"
                
                caption_run = caption_p.add_run(f"\nРисунок {fig_num} — {fig_title}")
                caption_run.font.name = 'Times New Roman'
                caption_run.font.size = docx.shared.Pt(11)
                caption_run.font.italic = True
                
                print(f"[DOCX ДИАГРАММА] Успешно вставили схему {image_path} под {heading_text_prefix}")
                return True
    return False

def populate_literature(doc, literature_list):
    for idx, p in enumerate(doc.paragraphs):
        if "Список литературы" in p.text:
            current_idx = idx + 1
            # Delete empty paragraphs immediately following it
            while current_idx < len(doc.paragraphs) and doc.paragraphs[current_idx].text.strip() == "" and not doc.paragraphs[current_idx].style.name.startswith("Heading"):
                p_to_del = doc.paragraphs[current_idx]
                p_to_del._element.getparent().remove(p_to_del._element)
                
            target_p = doc.paragraphs[idx]
            for lit_idx, item in enumerate(literature_list):
                lit_text = f"{lit_idx + 1}. {item}"
                
                new_p_el = docx.oxml.shared.OxmlElement('w:p')
                target_p._element.addnext(new_p_el)
                new_p = docx.text.paragraph.Paragraph(new_p_el, target_p._parent)
                new_p.style = 'Normal'
                
                # Format literature paragraph
                new_p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
                new_p.paragraph_format.line_spacing = 1.5
                new_p.paragraph_format.space_before = docx.shared.Pt(0)
                new_p.paragraph_format.space_after = docx.shared.Pt(4)
                
                run = new_p.add_run(lit_text)
                run.font.name = 'Times New Roman'
                run.font.size = docx.shared.Pt(12)
                
                # Force font XML
                rPr = run._r.get_or_add_rPr()
                rFonts = docx.oxml.shared.OxmlElement('w:rFonts')
                rFonts.set(docx.oxml.ns.qn('w:ascii'), 'Times New Roman')
                rFonts.set(docx.oxml.ns.qn('w:hAnsi'), 'Times New Roman')
                rFonts.set(docx.oxml.ns.qn('w:cs'), 'Times New Roman')
                rFonts.set(docx.oxml.ns.qn('w:eastAsia'), 'Times New Roman')
                rPr.append(rFonts)
                
                target_p = new_p
            return True
    return False

def adjust_tables_layout(doc):
    """Оптимизирует разметку таблиц: переводит секции с широкими таблицами в ландшафт,
    либо пропорционально масштабирует колонки, чтобы таблица не выходила за поля листа.
    """
    import copy
    
    def get_table_section_idx(doc, table_el):
        body = doc.element.body
        current_section_idx = 0
        for child in body:
            tag = child.tag.split('}')[-1]
            if tag == 'p':
                pPr = child.pPr
                if pPr is not None and pPr.xpath('./w:sectPr'):
                    current_section_idx += 1
            elif tag == 'tbl':
                if child == table_el:
                    return current_section_idx
        return current_section_idx

    def insert_section_break_before(doc, table_el, orig_sectPr):
        p = docx.oxml.shared.OxmlElement('w:p')
        pPr = docx.oxml.shared.OxmlElement('w:pPr')
        sectPr = copy.deepcopy(orig_sectPr)
        
        type_el = sectPr.xpath('./w:type')
        if not type_el:
            type_el = docx.oxml.shared.OxmlElement('w:type')
            sectPr.append(type_el)
        else:
            type_el = type_el[0]
        type_el.set(docx.oxml.ns.qn('w:val'), 'nextPage')
        
        pPr.append(sectPr)
        p.append(pPr)
        table_el.addprevious(p)
        return p

    def insert_section_break_after(doc, table_el, orig_sectPr):
        body = doc.element.body
        tbl_pos = list(body).index(table_el)
        if tbl_pos + 1 < len(body):
            next_sibling = list(body)[tbl_pos + 1]
            tag = next_sibling.tag.split('}')[-1]
            if tag == 'p':
                pPr = next_sibling.pPr
                if pPr is not None and pPr.xpath('./w:sectPr'):
                    return None
            elif tag == 'sectPr':
                return None
            
            p = docx.oxml.shared.OxmlElement('w:p')
            pPr = docx.oxml.shared.OxmlElement('w:pPr')
            sectPr = copy.deepcopy(orig_sectPr)
            
            type_el = sectPr.xpath('./w:type')
            if not type_el:
                type_el = docx.oxml.shared.OxmlElement('w:type')
                sectPr.append(type_el)
            else:
                type_el = type_el[0]
            type_el.set(docx.oxml.ns.qn('w:val'), 'nextPage')
            
            pPr.append(sectPr)
            p.append(pPr)
            next_sibling.addprevious(p)
            return p
        return None

    # Обрабатываем каждую таблицу по очереди
    for i, table in enumerate(doc.tables):
        if not table.rows:
            continue
            
        sec_idx = get_table_section_idx(doc, table._tbl)
        sec = doc.sections[sec_idx]
        
        # Вычисляем доступную ширину страницы (за вычетом полей)
        page_width = int(sec.page_width.twips)
        margin_left = int(sec.left_margin.twips) if sec.left_margin else 1440
        margin_right = int(sec.right_margin.twips) if sec.right_margin else 1440
        max_content_width = page_width - (margin_left + margin_right)
        
        # Вычисляем ширину таблицы (сумму ширин колонок первой строки)
        row0 = table.rows[0]
        total_width = 0
        for cell in row0.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.xpath('./w:tcW')
            w_val = 0
            if tcW:
                w_val = int(tcW[0].get(docx.oxml.ns.qn('w:w')) or 0)
            if w_val == 0 and cell.width:
                w_val = int(cell.width.twips)
            total_width += w_val
            
        print(f"[ТАБЛИЦА {i}] Раздел: {sec_idx}, Ориентация: {sec.orientation}, Ширина: {total_width} dxa, Лимит: {max_content_width} dxa")
        
        # Если таблица шире доступной области
        if total_width > max_content_width:
            num_cols = len(row0.cells)
            # Переводим в ландшафт только если в таблице >= 6 колонок (для больших таблиц)
            if sec.orientation != docx.enum.section.WD_ORIENT.LANDSCAPE and num_cols >= 6:
                print(f"[ТАБЛИЦА {i}] Ширина {total_width} превышает предел портрета {max_content_width} и колонок {num_cols} >= 6. Изолируем в новый раздел и переводим в LANDSCAPE...")
                
                orig_sectPr = sec._sectPr
                insert_section_break_before(doc, table._tbl, orig_sectPr)
                insert_section_break_after(doc, table._tbl, orig_sectPr)
                
                # Ищем новый индекс раздела для таблицы после добавления разрывов
                sec_idx = get_table_section_idx(doc, table._tbl)
                sec = doc.sections[sec_idx]
                
                # Переводим раздел в альбомную ориентацию
                sec.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE
                if sec.page_width < sec.page_height:
                    w, h = sec.page_width, sec.page_height
                    sec.page_width = h
                    sec.page_height = w
                
                # Поворачиваем margins
                margin_top = sec.top_margin
                margin_bottom = sec.bottom_margin
                margin_left = sec.left_margin
                margin_right = sec.right_margin
                
                sec.top_margin = margin_right
                sec.bottom_margin = margin_left
                sec.left_margin = margin_top
                sec.right_margin = margin_bottom
                
                # Пересчитываем новый лимит ширины контента для ландшафтного листа
                page_width = int(sec.page_width.twips)
                margin_left_val = int(sec.left_margin.twips) if sec.left_margin else 1440
                margin_right_val = int(sec.right_margin.twips) if sec.right_margin else 1440
                max_content_width = page_width - (margin_left_val + margin_right_val)
                
            # Если таблица всё еще шире лимита контента (даже в ландшафте), масштабируем ячейки пропорционально
            if total_width > max_content_width and total_width > 0:
                scale = max_content_width / total_width
                print(f"[ТАБЛИЦА {i}] Превышает предел даже в ландшафте. Масштабируем ячейки с коэффициентом {scale:.2f}")
                for r in table.rows:
                    for cell in r.cells:
                        tcPr = cell._tc.get_or_add_tcPr()
                        tcW = tcPr.xpath('./w:tcW')
                        if tcW:
                            w_val = int(tcW[0].get(docx.oxml.ns.qn('w:w')) or 0)
                            if w_val == 0 and cell.width:
                                w_val = int(cell.width.twips)
                            new_w = int(w_val * scale)
                            tcW[0].set(docx.oxml.ns.qn('w:w'), str(new_w))
                        else:
                            # Если элемента w:tcW нет, создаем его с масштабированной шириной
                            w_val = int(cell.width.twips) if cell.width else int(total_width / len(r.cells))
                            new_w = int(w_val * scale)
                            tcW_el = docx.oxml.shared.OxmlElement('w:tcW')
                            tcW_el.set(docx.oxml.ns.qn('w:type'), 'dxa')
                            tcW_el.set(docx.oxml.ns.qn('w:w'), str(new_w))
                            tcPr.append(tcW_el)

def build_word_report(json_data_path, template_path, output_filename):
    """Loads JSON data generated by agents, populates templaredoc.docx, and saves it."""
    print(f"[DOCX] Loading data from {json_data_path}...")
    with open(json_data_path, "r", encoding="utf-8") as f:
        report_data = json.load(f)
        
    print(f"[DOCX] Opening template {template_path}...")
    doc = docx.Document(template_path)
    
    # Force Times New Roman on all existing cells in all tables (headers, default texts)
    print("[DOCX] Forcing Times New Roman on all table text...")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.name = 'Times New Roman'
                        rPr = r._r.get_or_add_rPr()
                        rFonts = docx.oxml.shared.OxmlElement('w:rFonts')
                        rFonts.set(docx.oxml.ns.qn('w:ascii'), 'Times New Roman')
                        rFonts.set(docx.oxml.ns.qn('w:hAnsi'), 'Times New Roman')
                        rFonts.set(docx.oxml.ns.qn('w:cs'), 'Times New Roman')
                        rFonts.set(docx.oxml.ns.qn('w:eastAsia'), 'Times New Roman')
                        rPr.append(rFonts)
    
    # 1. Update Title Page
    print("[DOCX] Filling Title Page...")
    topic = report_data.get('topic', '')
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt.startswith("ТЕМА ПРОЕКТА:"):
            p.text = ""
            r1 = p.add_run("ТЕМА ПРОЕКТА:\n")
            r1.font.name = 'Times New Roman'
            r1.font.size = docx.shared.Pt(12)
            r2 = p.add_run(topic.upper())
            r2.bold = True
            r2.font.name = 'Times New Roman'
            r2.font.size = docx.shared.Pt(14)
        elif txt.startswith("Слушатель программы:"):
            p.text = ""
            r1 = p.add_run("Слушатель программы: ")
            r1.font.name = 'Times New Roman'
            r1.font.size = docx.shared.Pt(12)
            r2 = p.add_run("_______________________")
            r2.bold = True
            r2.font.name = 'Times New Roman'
            r2.font.size = docx.shared.Pt(12)
        elif txt.startswith("Руководитель проектной работы:"):
            p.text = ""
            r1 = p.add_run("Руководитель проектной работы: ")
            r1.font.name = 'Times New Roman'
            r1.font.size = docx.shared.Pt(12)
            r2 = p.add_run("_______________________")
            r2.bold = True
            r2.font.name = 'Times New Roman'
            r2.font.size = docx.shared.Pt(12)
        elif txt.startswith("Наставник:"):
            p.text = ""
            r1 = p.add_run("Наставник: ")
            r1.font.name = 'Times New Roman'
            r1.font.size = docx.shared.Pt(12)
            r2 = p.add_run("_______________________")
            r2.bold = True
            r2.font.name = 'Times New Roman'
            r2.font.size = docx.shared.Pt(12)
            
    # Оптимизация титульного листа: убираем лишние пустые абзацы, чтобы 'Севастополь 2026' не уезжал на стр. 2
    sev_idx = -1
    ann_idx = -1
    for idx, p in enumerate(doc.paragraphs):
        txt_lower = p.text.lower().strip()
        if "севастополь" in txt_lower and ("2026" in txt_lower or "2025" in txt_lower):
            sev_idx = idx
        elif "аннотация проекта" in txt_lower and p.style.name.startswith("Heading"):
            ann_idx = idx
            
    if sev_idx != -1:
        # 1. Принудительный разрыв страницы перед Аннотацией
        if ann_idx != -1:
            doc.paragraphs[ann_idx].paragraph_format.page_break_before = True
            # Удаляем все пустые абзацы между 'Севастополь' и 'Аннотация'
            for idx in range(ann_idx - 1, sev_idx, -1):
                p_between = doc.paragraphs[idx]
                if p_between.text.strip() == "":
                    p_between._element.getparent().remove(p_between._element)
                    
        # 1.5. Принудительный разрыв страницы перед Оглавлением (TOC)
        from docx.oxml.text.paragraph import CT_P
        from docx.text.paragraph import Paragraph
        for el in doc.element.body.iter():
            if isinstance(el, CT_P):
                p = Paragraph(el, doc)
                if p.text.strip().lower() == "оглавление":
                    p.paragraph_format.page_break_before = True
                    print("[DOCX] Найдено оглавление, установлен разрыв страницы перед ним.")
                    break
        
        # Пересчитываем индекс Севастополя после удаления абзацев
        for idx, p in enumerate(doc.paragraphs):
            txt_lower = p.text.lower().strip()
            if "севастополь" in txt_lower and ("2026" in txt_lower or "2025" in txt_lower):
                sev_idx = idx
                break
                
        # 2. Сокращаем пустые абзацы перед 'Севастополь'
        empty_before = []
        curr = sev_idx - 1
        while curr >= 0:
            p_before = doc.paragraphs[curr]
            if p_before.text.strip() == "":
                empty_before.append(p_before)
            else:
                break
            curr -= 1
            
        keep_count = 2
        if len(empty_before) > keep_count:
            for p_del in empty_before[keep_count:]:
                p_del._element.getparent().remove(p_del._element)
            print(f"[DOCX] Титульный лист: оптимизирована разметка, удалено {len(empty_before) - keep_count} пустых абзацев.")

    # 2. Fill Annotation (Table 0)
    print("[DOCX] Filling Annotation Table...")
    if len(doc.tables) > 0:
        populate_annotation_table(doc.tables[0], report_data)
    print("[DOCX] Filling text sections...")
    replace_text_under_heading(doc, "1. Масштаб", report_data.get('section_1_scale', ''))
    replace_text_under_heading(doc, "2. Обоснование", report_data.get('section_2_problems', ''))
    replace_text_under_heading(doc, "3. Цель", report_data.get('section_3_goal_tasks', ''))
    replace_text_under_heading(doc, "4. Обоснование новизны", report_data.get('section_4_novelty', ''))
    replace_text_under_heading(doc, "5. Дорожная карта", f"Ниже представлен подробный календарный план и дорожная карта реализации проекта по теме «{topic}», включая ключевые этапы, мероприятия, сроки и ответственных лиц.")
    replace_text_under_heading(doc, "6. Результаты", report_data.get('section_6_results', ''))
    replace_text_under_heading(doc, "7. Риски", f"В данном разделе приведен всесторонний анализ рисков проекта по теме «{topic}», оценка их вероятности, потенциального влияния, а также разработанные превентивные меры и планы реагирования.")
    
    # 3.5. Fill Appendices 1-4
    print("[DOCX] Filling Appendices 1-4...")
    replace_text_under_appendix(doc, "Приложение 1", report_data.get('appendix_1_problems', ''))
    replace_text_under_appendix(doc, "Приложение 2", report_data.get('appendix_2_solutions', ''))
    replace_text_under_appendix(doc, "Приложение 3", report_data.get('appendix_3_practices', ''))
    replace_text_under_appendix(doc, "Приложение 4", report_data.get('appendix_4_target_groups', ''))
    
    # Генерация диаграмм с помощью локального Pillow-генератора на основе Gemini/парсинга промпта
    topic = report_data.get('topic', '')
    safe_topic_name = "".join([c if c.isalnum() or c in (' ', '_') else '_' for c in topic])
    safe_topic_name = safe_topic_name.strip().replace(" ", "_")[:50]
    
    try:
        from src.diagram_generator import generate_report_diagrams
        diagram_1_path, diagram_2_path = generate_report_diagrams(json_data_path, safe_topic_name)
    except Exception as diag_err:
        print(f"[DOCX] Ошибка при импорте или запуске генератора диаграмм: {diag_err}")
        # Попытка найти резервные файлы
        diagram_1_path = os.path.join("output", f"diagram_app1_{safe_topic_name}.png")
        diagram_2_path = os.path.join("output", f"diagram_app2_{safe_topic_name}.png")
        
    insert_diagram_to_docx(doc, "Приложение 1", diagram_1_path)
    insert_diagram_to_docx(doc, "Приложение 2", diagram_2_path)
    
    # 4. Fill Table 1 (Roadmap)
    print("[DOCX] Rebuilding Calendar Table...")
    if len(doc.tables) > 1:
        rebuild_roadmap_table(doc.tables[1], report_data.get('roadmap', []))
        
    # 5. Fill Table 2 (Risks)
    print("[DOCX] Rebuilding Risks Table...")
    if len(doc.tables) > 2:
        rebuild_risks_table(doc.tables[2], report_data.get('risks', []))
        
    # 6. Fill Literature
    print("[DOCX] Populating literature list...")
    populate_literature(doc, report_data.get('literature', []))
    
    # 7. Fill Table 3 (Stakeholders)
    print("[DOCX] Rebuilding Stakeholders Table...")
    if len(doc.tables) > 3:
        rebuild_stakeholders_table(doc.tables[3], report_data.get('stakeholders', []))
        
    # 8. Fill Table 4 (Communication Plan)
    print("[DOCX] Rebuilding Communication Plan Table...")
    if len(doc.tables) > 4:
        rebuild_communication_table(doc.tables[4], report_data.get('communication_plan', []))
        
    # Оптимизация разметки таблиц
    print("[DOCX] Adjusting tables layout to fit the page and flipping orientation if needed...")
    adjust_tables_layout(doc)
        
    # Разделение разделов с новой страницы (ГОСТ/РАНХиГС)
    print("[DOCX] Adding page breaks before major headings...")
    for idx, p in enumerate(doc.paragraphs):
        if p.style.name.startswith("Heading 1"):
            txt = p.text.strip().lower()
            # Добавляем разрыв страницы перед всеми numbered разделами и приложениями
            if any(txt.startswith(prefix) for prefix in ["1.", "2.", "3.", "4.", "5.", "6.", "7.", "приложение"]) or "литератур" in txt:
                # Проверяем, не идет ли перед этим заголовком разрыв раздела (sectPr)
                is_first_in_section = False
                if idx > 0:
                    prev_p = doc.paragraphs[idx - 1]
                    prev_pPr = prev_p._element.pPr
                    if prev_pPr is not None and len(prev_pPr.xpath('./w:sectPr')) > 0:
                        is_first_in_section = True
                
                # Также проверим, не пустой ли предыдущий абзац и не содержит ли он разрыв раздела
                if idx > 1 and not is_first_in_section:
                    prev_prev_p = doc.paragraphs[idx - 2]
                    prev_prev_pPr = prev_prev_p._element.pPr
                    if prev_prev_pPr is not None and len(prev_prev_pPr.xpath('./w:sectPr')) > 0:
                        if doc.paragraphs[idx - 1].text.strip() == "":
                            is_first_in_section = True
                
                # Задаем разрыв страницы только если заголовок не является первым в новой секции
                if not is_first_in_section:
                    p.paragraph_format.page_break_before = True

    # Глобальная очистка пустых абзацев после титульного листа (предотвращает появление пустых страниц)
    print("[DOCX] Global cleanup of empty paragraphs to prevent blank pages...")
    first_heading_idx = -1
    for idx, p in enumerate(doc.paragraphs):
        if p.style.name.startswith("Heading 1"):
            first_heading_idx = idx
            break
            
    if first_heading_idx != -1:
        # Идем с конца в начало, чтобы безопасно удалять элементы по индексу
        for idx in range(len(doc.paragraphs) - 1, first_heading_idx, -1):
            p = doc.paragraphs[idx]
            # Если текст пустой, в абзаце нет встроенных картинок/рисунков и это не разрыв раздела w:sectPr
            if p.text.strip() == "" and len(p._element.xpath('.//w:drawing')) == 0 and len(p._element.xpath('.//w:sectPr')) == 0:
                p._element.getparent().remove(p._element)
        
    # Настройка автоматического обновления оглавления (TOC) при открытии документа в Word
    print("[DOCX] Enabling automatic field update (TOC page numbers) on document open...")
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        settings = doc.settings.element
        update_fields = settings.find(qn('w:updateFields'))
        if update_fields is None:
            update_fields = OxmlElement('w:updateFields')
            update_fields.set(qn('w:val'), 'true')
            settings.append(update_fields)
        else:
            update_fields.set(qn('w:val'), 'true')
    except Exception as e:
        print(f"[DOCX] Не удалось включить автообновление полей: {e}")
        
    # Save the document
    output_path = os.path.join("output", output_filename)
    print(f"[DOCX] Saving to {output_path}...")
    doc.save(output_path)
    print("[DOCX] Saved successfully.")
    return output_path
